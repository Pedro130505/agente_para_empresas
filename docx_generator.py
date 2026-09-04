import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from pathlib import Path
from config import TEMPLATE_DOCX_PATH, OUTPUT_DIR

COLOR_HEADER = RGBColor(0x00, 0x33, 0x66)    # Azul Escuro Institucional UFMG
COLOR_SUBHEADER = RGBColor(0x00, 0x66, 0x99) # Azul Secundário
COLOR_TEXT = RGBColor(0x2B, 0x2B, 0x2B)      # Grafite Escuro

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(17)
    run.font.bold = True
    run.font.color.rgb = COLOR_HEADER
    return p

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = COLOR_SUBHEADER
    return p

def add_section_header(doc, number_str, title_str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    r_num = p.add_run(number_str + " ")
    r_num.font.name = 'Calibri'
    r_num.font.size = Pt(12)
    r_num.font.bold = True
    r_num.font.color.rgb = COLOR_SUBHEADER
    
    r_title = p.add_run(title_str.upper())
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(12)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_HEADER
    return p

def add_body_p(doc, text, bold_prefix=None):
    if not text:
        return None
    paragraphs = text.split("\n\n")
    last_p = None
    for i, part in enumerate(paragraphs):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix and i == 0:
            r_pre = p.add_run(bold_prefix)
            r_pre.font.name = 'Calibri'
            r_pre.font.size = Pt(10)
            r_pre.font.bold = True
            r_pre.font.color.rgb = COLOR_TEXT
            
        r_text = p.add_run(part)
        r_text.font.name = 'Calibri'
        r_text.font.size = Pt(10)
        r_text.font.color.rgb = COLOR_TEXT
        last_p = p
    return last_p

def add_callout(doc, title, text, bg_hex="F0F4F8", border_hex="003366", icon="📌"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r_t = p.add_run(f"{icon} {title}\n")
    r_t.font.name = 'Calibri'
    r_t.font.size = Pt(10)
    r_t.font.bold = True
    r_t.font.color.rgb = COLOR_HEADER
    
    r_b = p.add_run(text)
    r_b.font.name = 'Calibri'
    r_b.font.size = Pt(9.5)
    r_b.font.italic = True
    r_b.font.color.rgb = COLOR_TEXT
    
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(4)

def add_fairs_table(doc, fairs_data):
    """Gera a tabela estruturada de Inteligência Competitiva de Feiras (WI, RC, UFRJ, PUC)."""
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    hdr_cells = tbl.rows[0].cells
    hdr_cells[0].width = Inches(2.2)
    hdr_cells[1].width = Inches(0.9)
    hdr_cells[2].width = Inches(0.9)
    hdr_cells[3].width = Inches(2.5)
    
    for c in hdr_cells:
        set_cell_background(c, "003366")
        set_cell_margins(c, top=80, bottom=80, left=100, right=100)
        
    p0 = hdr_cells[0].paragraphs[0]
    p0.add_run("Feira / Evento").font.bold = True
    p0.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p0.runs[0].font.size = Pt(9)
    
    p1 = hdr_cells[1].paragraphs[0]
    p1.add_run("2025").font.bold = True
    p1.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p1.runs[0].font.size = Pt(9)
    
    p2 = hdr_cells[2].paragraphs[0]
    p2.add_run("2026").font.bold = True
    p2.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p2.runs[0].font.size = Pt(9)

    p3 = hdr_cells[3].paragraphs[0]
    p3.add_run("Detalhes da Participação").font.bold = True
    p3.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p3.runs[0].font.size = Pt(9)

    for item in fairs_data:
        row_cells = tbl.add_row().cells
        row_cells[0].width = Inches(2.2)
        row_cells[1].width = Inches(0.9)
        row_cells[2].width = Inches(0.9)
        row_cells[3].width = Inches(2.5)
        
        set_cell_background(row_cells[0], "F9FAFB")
        set_cell_background(row_cells[1], "FFFFFF")
        set_cell_background(row_cells[2], "FFFFFF")
        set_cell_background(row_cells[3], "FFFFFF")
        
        for c in row_cells:
            set_cell_margins(c, top=60, bottom=60, left=100, right=100)
            
        p_f = row_cells[0].paragraphs[0]
        r_f = p_f.add_run(item.get("feira", ""))
        r_f.font.bold = True
        r_f.font.size = Pt(8.5)
        
        p_s25 = row_cells[1].paragraphs[0]
        r_s25 = p_s25.add_run(item.get("status_2025", "Não"))
        r_s25.font.size = Pt(8.5)
        r_s25.font.bold = True
        if str(item.get("status_2025")).lower() in ["sim", "true"]:
            r_s25.font.color.rgb = RGBColor(0x28, 0xA7, 0x45)
            
        p_s26 = row_cells[2].paragraphs[0]
        r_s26 = p_s26.add_run(item.get("status_2026", "Não"))
        r_s26.font.size = Pt(8.5)
        r_s26.font.bold = True
        if str(item.get("status_2026")).lower() in ["sim", "true"]:
            r_s26.font.color.rgb = RGBColor(0x28, 0xA7, 0x45)

        p_det = row_cells[3].paragraphs[0]
        r_det = p_det.add_run(item.get("detalhes", ""))
        r_det.font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_objections_table(doc, objections_list):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    hdr_cells = tbl.rows[0].cells
    hdr_cells[0].width = Inches(2.3)
    hdr_cells[1].width = Inches(4.2)
    set_cell_background(hdr_cells[0], "003366")
    set_cell_background(hdr_cells[1], "003366")
    set_cell_margins(hdr_cells[0], top=100, bottom=100, left=120, right=120)
    set_cell_margins(hdr_cells[1], top=100, bottom=100, left=120, right=120)
    
    p0 = hdr_cells[0].paragraphs[0]
    r0 = p0.add_run("Objeção Provável do Cliente")
    r0.font.bold = True
    r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r0.font.size = Pt(9.5)
    
    p1 = hdr_cells[1].paragraphs[0]
    r1 = p1.add_run("Resposta Estruturada & Contra-Argumento Tático")
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r1.font.size = Pt(9.5)

    for item in objections_list:
        row_cells = tbl.add_row().cells
        row_cells[0].width = Inches(2.3)
        row_cells[1].width = Inches(4.2)
        
        set_cell_background(row_cells[0], "F9FAFB")
        set_cell_background(row_cells[1], "FFFFFF")
        set_cell_margins(row_cells[0], top=80, bottom=80, left=120, right=120)
        set_cell_margins(row_cells[1], top=80, bottom=80, left=120, right=120)
        
        p_obj = row_cells[0].paragraphs[0]
        p_obj.paragraph_format.space_after = Pt(0)
        r_obj = p_obj.add_run(item.get("objecao", ""))
        r_obj.font.bold = True
        r_obj.font.size = Pt(9)
        r_obj.font.color.rgb = COLOR_TEXT
        
        p_resp = row_cells[1].paragraphs[0]
        p_resp.paragraph_format.space_after = Pt(0)
        r_resp = p_resp.add_run(item.get("resposta", ""))
        r_resp.font.size = Pt(9)
        r_resp.font.color.rgb = COLOR_TEXT

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(6)

def generate_one_page_docx(company_data, ai_analysis, output_dir=OUTPUT_DIR, template_path=TEMPLATE_DOCX_PATH):
    nome = company_data["nome"]
    
    if Path(template_path).exists():
        doc = docx.Document(template_path)
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)
        for t in list(doc.tables):
            t._element.getparent().remove(t._element)
    else:
        doc = docx.Document()

    # Cabeçalho Principal
    add_title(doc, f"DOSSIÊ ESTRATÉGICO DE INTELIGÊNCIA: {nome.upper()}")
    add_subtitle(doc, "Manual de Preparação e Suporte para Reunião Comercial — UFMG Hub / Feira de Carreiras")

    # Tabela 0: Ficha Técnica de Inteligência
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    atuacao_raw = ai_analysis.get("atuacao_bh_mg_detalhada", "")
    tem_mg = "Sim" if "MINAS GERAIS: SIM" in atuacao_raw.upper() else ("Sim" if "MG" in atuacao_raw else "A confirmar")
    tem_bh = "Sim" if "BELO HORIZONTE: SIM" in atuacao_raw.upper() else ("Sim" if "BELO HORIZONTE" in atuacao_raw.upper() else "A confirmar")

    labels_data = [
        ("Empresa Prospectada", nome),
        ("Presença Regional (MG / BH)", f"Minas Gerais: {tem_mg} | Belo Horizonte: {tem_bh}"),
        ("Histórico Feira UFMG", f"2024: {company_data['participou_2024']} (Cota: {company_data['cota_2024']}) | 2025: {company_data['participou_2025']} (Cota: {company_data['cota_2025']}) | 2026: {company_data['participou_2026']} (Cota: {company_data['cota_2026']})"),
        ("Contato / Responsável", f"{company_data.get('nome_contato') or company_data.get('responsavel_2026') or 'A mapear'} ({company_data.get('email') or 'Sem e-mail cadastrado'})"),
        ("Feiras Mapeadas", "Workshop Integrativo (WI) e PUC Carreiras")
    ]

    for idx, (label, val) in enumerate(labels_data):
        row = table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(4.3)
        
        set_cell_background(cell_lbl, "EBF3FA")
        set_cell_background(cell_val, "FAFAFA")
        set_cell_margins(cell_lbl, top=70, bottom=70, left=100, right=100)
        set_cell_margins(cell_val, top=70, bottom=70, left=100, right=100)
        
        p0 = cell_lbl.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(9)
        r0.font.color.rgb = COLOR_HEADER
        
        p1 = cell_val.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(val)
        r1.font.size = Pt(9)
        r1.font.color.rgb = COLOR_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 1. Visão Geral & Atuação no Mercado
    add_section_header(doc, "1.0", "Visão Geral Institucional & Foco de Mercado (Brasil & MG)")
    add_body_p(doc, ai_analysis.get("resumo_extenso", ""))

    # 2. Atuação em Belo Horizonte e MG
    add_section_header(doc, "2.0", "Presença Física, Fábricas e Localização de Operações em BH & MG")
    add_body_p(doc, ai_analysis.get("atuacao_bh_mg_detalhada", ""))

    # 3. Histórico de Relacionamento e Patrocínios
    add_section_header(doc, "3.0", "Histórico de Relacionamento & Análise de Patrocínios (2024 - 2026)")
    add_body_p(doc, ai_analysis.get("historico_relacionamento_analise", ""))

    # 4. Programas de Estágio e Trainee
    add_section_header(doc, "4.0", "Programas de Atração de Talentos em BH & MG (Estágio & Trainee)")
    add_body_p(doc, ai_analysis.get("programas_estagio_trainee_completo", ""))

    # 5. Inteligência Competitiva e Outras Feiras (Tabela Concreta)
    add_section_header(doc, "5.0", "Inteligência Competitiva & Presença Confirmada em Outros Eventos")
    add_fairs_table(doc, ai_analysis.get("outras_feiras_tabela", []))

    # 6. Posicionamento de Marca & ESG
    add_section_header(doc, "6.0", "Posicionamento de Marca, ESG & Inovação")
    add_body_p(doc, ai_analysis.get("posicionamento_esg_inovacao", ""))

    # 7. Manual de Reunião de Vendas (Guia Tático)
    add_section_header(doc, "7.0", "Guia Tático para Reunião Comercial com o Cliente")
    
    # 7.1 Ganchos de Abertura
    ganchos_txt = "\n".join([f"• {g}" for g in ai_analysis.get("guia_reuniao_ganchos", [])])
    add_callout(doc, "Ganchos de Abertura para Iniciar a Reunião", ganchos_txt, bg_hex="F4F6F9", border_hex="003366", icon="💡")

    # 7.2 Pitch Principal
    add_callout(doc, "Discurso Principal de Valor (Pitch B2B)", ai_analysis.get("guia_reuniao_pitch", ""), bg_hex="FFFDF0", border_hex="B8860B", icon="🎯")

    # 7.3 Tabela de Objeções
    add_body_p(doc, "Abaixo estão as objeções mais prováveis do cliente durante a reunião e as respostas recomendadas:", bold_prefix="Matriz de Objeções & Respostas Táticas: ")
    add_objections_table(doc, ai_analysis.get("guia_reuniao_objecoes", []))

    clean_filename = "".join(c for c in nome if c.isalnum() or c in (' ', '_', '-')).strip()
    output_filename = Path(output_dir) / f"Dossie_Estrategico_{clean_filename}.docx"
    doc.save(output_filename)
    return output_filename
